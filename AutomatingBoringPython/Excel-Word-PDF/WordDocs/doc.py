import docx

doc = docx.Document(r'C:\Users\demo.docx') #return new doc object

doc.paragraphs #paragraphs variable that is a list of paragraph objects

doc.paragraphs[0].text #extract text from the first paragraph object

p = doc.paragraphs[1]

p.runs #this is a list of run object, a new run start when there is a change in doc style eg. regular text, bold text, italic text etc

p.runs[0].text #extracting run object text

p.runs[1].bold #return true b/c the run have a bold object

p.runs[0].bold == None #unbold somehting

p.runs[3].underline == True #making underline for that runs

p.runs[3].text = 'italic and underline' #changing that run's text

doc.save(r'C:\Users\demo2.docx') #saving the doc

p.style # See what style it have

p.style = 'Title' #changing the style to title style

doc.save(r'C:\Users\demo3.docx')

'''
Creating New Doc
'''

newDoc = docx.Document()
newDoc.add_paragraph('Hello World!') #adding new paragraph and add what text it should have
newDoc.save(r'C:\Users\demo4.docx')

p1 = newDoc.paragraphs[0]
p1.add_run('This is a new Run') #this add new run to the first paragraph
